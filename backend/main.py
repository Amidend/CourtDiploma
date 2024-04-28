import datetime
import io
from urllib import response

from fastapi import FastAPI, File, UploadFile, Form, Depends, APIRouter, HTTPException
from typing import List
from io import BytesIO
from docx import Document
from fastapi.responses import StreamingResponse
from fastapi.params import Body
from urllib.parse import quote

from pydantic import BaseModel
from sqlalchemy.exc import SQLAlchemyError, IntegrityError
from sqlalchemy.orm import Session
from fastapi.middleware.cors import CORSMiddleware
from starlette import status
from starlette.responses import JSONResponse, FileResponse

from backend.aut.UserCreate import UserCreate
from backend.aut.aut import authenticate_user, get_admin_user
from backend.models import models
from backend.routers import judges, companies, templates, users, documents
from models.database import get_db, create_tables  # Import database session
from models.models import Template, Judge, Company, User  # Import SQLAlchemy models
from fastapi import APIRouter, Depends
from fastapi.security import OAuth2PasswordRequestForm
from fastapi import APIRouter, Depends, Response


app = FastAPI()

from backend.models.models import Judge
from backend.models.database import get_db

router = APIRouter()


origins = [
    "http://127.0.0.1:5173",
    "http://localhost:5173",
    "http://localhost",
    "http://localhost:8080",

]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

create_tables()

@app.post("/admin/upload_template")
async def upload_template(file: UploadFile = File(...), db: Session = Depends(get_db)):
    content = await file.read()  # Read the binary content
    template = Template(name=file.filename, content=content)
    db.add(template)
    db.commit()
    db.refresh(template)
    return {"message": f"Template '{file.filename}' uploaded successfully with ID: {template.id}"}

class GenerateDocRequest(BaseModel):
    template_name: int
    data: List[str]
# СайтЭноном.docx
@app.post("/user/generate_document")
async def generate_document(
        request_data: GenerateDocRequest = Body(...),
        db: Session = Depends(get_db),
):
    # 1. Retrieve the template content
    template = db.query(Template).filter_by(id=request_data.template_name).first()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    document = Document(BytesIO(template.content))

    # 3. Replace placeholders with data (assuming placeholders are like {0}, {1}, etc.)
    # Разбиваем его на части с помощью запятой
    parts = request_data.data

    data = [db.query(Judge).filter_by(id=request_data.data[0]).first().fio,db.query(Company).filter_by(id=request_data.data[1]).first().name]
    print(data)
    for i, value in enumerate(data):
        for paragraph in document.paragraphs:
            paragraph.text = paragraph.text.replace("{" + str(i) + "}", value)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = cell.text.replace("{" + str(i) + "}", value)

    # 4. Save the document to a BytesIO object
    new_file = BytesIO()
    document.save(new_file)
    new_file.seek(0)






    document_record = models.Document(
        name=f"{request_data.template_name}_generated",
        description="Generated document",
        type="docx",
        date=datetime.date.today(),
        judge_id = parts[0],
        company_id = parts[1],
        file=new_file.getvalue(),  # Store the file content directly
    )
    db.add(document_record)
    db.commit()
    db.refresh(document_record)

    encoded_filename = quote(f"{request_data.template_name}_generated.docx")

    # Return the file as a streaming response
    return StreamingResponse(
        new_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
        },
    )

@router.get("/options")
async def get_selection_options(db: Session = Depends(get_db)):
    templates = db.query(Template.id, Template.name).all()
    judges = db.query(Judge.id, Judge.fio).all()
    companies = db.query(Company.id, Company.name).all()

    print("Templates:", templates)  # Add print statement for debugging
    print("Judges:", judges)
    print("Companies:", companies)

    return {
        "templates": [{"id": t[0], "name": t[1]} for t in templates],
        "judges": [{"id": j[0], "fio": j[1]} for j in judges],
        "companies": [{"id": c[0], "name": c[1]} for c in companies],
    }



#aut
@router.post("/users", dependencies=[Depends(get_admin_user)])
async def register_user(user: UserCreate, db: Session = Depends(get_db)):
    try:
        pass
        return
    except IntegrityError as e:
        # Handle username uniqueness constraint violation
        db.rollback()
        raise HTTPException(status_code=400, detail="Username already exists")
    except SQLAlchemyError as e:
        # Handle other SQLAlchemy errors
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Database error: {e}")

@router.post("/login")
async def login(response: Response, db: Session = Depends(get_db), data: dict = Body(...)):
    try:
        username = data.get("username")
        password = data.get("password")

        if not username or not password:
            raise HTTPException(status_code=422, detail="Username and password are required")

        user = authenticate_user(db, username, password)
        if not user:
            raise HTTPException(status_code=401, detail="Invalid credentials")

        return JSONResponse(status_code=200, content={"userId": user.id, "role": user.role})
    except HTTPException as e:
        raise e



@router.post("/logout")
async def logout(response: Response):
    response.delete_cookie(key="user_id")
    return JSONResponse(content={"message": "Logout successful"})






app.include_router(judges.router)

app.include_router(router)
app.include_router(documents.router)
app.include_router(companies.router)
app.include_router(templates.router)
app.include_router(users.router)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, port=8000)


